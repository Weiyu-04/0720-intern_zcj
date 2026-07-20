#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
用模板生成《事故后评估》表2、表3。默认产出 4 个文件：
  - 表2/表3 各一个 Word(.docx)（官方模板，保留勾选占位/填表人行）
  - 表2/表3 各一个 Excel(.xlsx)（openpyxl 现绘，忠实镜像官方版式，供用户使用）
docx 与 xlsx 共用同一份输入 JSON，单一事实来源。

用法：
    python generate_tables.py <input.json> <output_dir>

输入 JSON 结构：
{
  "table2": [
    {"no":"1.1","requirement":"整改要求原文","subject":"整改主体",
     "content":"2.6.1 危险作业管理制度；2.6.2 危大工程管理制度","method":"资料审查、座谈问询、实地检查、专业评查",
     "materials":"建议提供的材料...","note":""}
  ],
  "table3_groups": [
    {"group_title": null, "rows":[
      {"no":"1","person":"责任人/单位（隶属单位）","suggestion":"处理建议",
       "provider":"提供材料单位","materials":"需提供材料","note":""}
    ]},
    {"group_title": "企业内部处理情况", "rows":[ ... ]}
  ]
}
说明：模板 templates/表2模板.docx、表3模板.docx 提供表头/列宽/边框/填表人行以及
「整改情况(符合/基本符合/不符合)」「处理情况(符合/不符合)」占位——这两列保持模板原样，专家勾选。
"""
import sys
import json
import os
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# JSON 加载 + 可读报错：报告正文常带引号，模型手写合并 JSON 时易把它写成未转义的半角 "，
# 直接 json.load 只抛字节偏移的 JSONDecodeError（模型看不懂→反复 replace/sed 打地鼠）。
# load_json_or_explain 严格解析，失败则给"定位到行 + 一次性正确做法"的处方式报错（不静默改文本）。
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from robust_json import load_json_or_explain

SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TPL2 = os.path.join(SKILL_DIR, "templates", "表2模板.docx")
TPL3 = os.path.join(SKILL_DIR, "templates", "表3模板.docx")


def _write_cell(cell, text):
    """清空单元格并写入文本，保持宋体（保留模板的列宽/边框等表级格式）。"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    run.font.size = Pt(9)
    run.font.name = "宋体"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), "宋体")


def _clone_data_row(table):
    """以模板的数据行(第2行, index=1)为原型克隆一行, 插到填表人行(末行)之前, 返回新行。"""
    proto_tr = table.rows[1]._tr
    footer_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(proto_tr)
    footer_tr.addprevious(new_tr)
    # 重新按 XML 顺序找到刚插入的行对象
    for r in table.rows:
        if r._tr is new_tr:
            return r
    return table.rows[-2]


def _remove_prototype(table):
    """删除模板自带的空原型数据行(填充完之后)。"""
    proto_tr = table.rows[1]._tr
    proto_tr.getparent().remove(proto_tr)


def build_table2(records, outpath):
    doc = Document(TPL2)
    table = doc.tables[0]
    for rec in records:
        row = _clone_data_row(table)
        vals = [rec.get("no", ""), rec.get("requirement", ""), rec.get("subject", ""),
                rec.get("content", ""), rec.get("method", ""), rec.get("materials", "")]
        for i, v in enumerate(vals):      # 列0-5填写；列6(整改情况)/列7(说明)保留模板占位
            _write_cell(row.cells[i], v)
        # 章节标题行((一)(二)(三)，其余列填"-")：整改情况也填"-"，不显示勾选项
        if rec.get("subject", "").strip() == "-":
            _write_cell(row.cells[6], "-")
        if rec.get("note"):
            _write_cell(row.cells[7], rec["note"])
    _remove_prototype(table)
    doc.save(outpath)


def build_table3(groups, outpath):
    doc = Document(TPL3)
    table = doc.tables[0]
    ncol = len(table.columns)
    for g in groups:
        gt = g.get("group_title")
        if gt:
            row = _clone_data_row(table)
            merged = row.cells[0]
            for i in range(1, ncol):
                merged = merged.merge(row.cells[i])
            _write_cell(merged, gt)
            merged.paragraphs[0].runs[0].bold = True
        for rec in g.get("rows", []):
            row = _clone_data_row(table)
            vals = [rec.get("no", ""), rec.get("person", ""), rec.get("suggestion", ""),
                    rec.get("provider", ""), rec.get("materials", "")]
            for i, v in enumerate(vals):  # 列0-4填写；列5(处理情况)/列6(说明)保留模板占位
                _write_cell(row.cells[i], v)
            # 组标题行(如"行刑衔接情况")：其余列留空，不显示勾选项
            if not (rec.get("suggestion") or rec.get("provider") or rec.get("materials")):
                _write_cell(row.cells[5], "")
                row.cells[1].paragraphs[0].runs[0].bold = True
            if rec.get("note"):
                _write_cell(row.cells[6], rec["note"])
    _remove_prototype(table)
    doc.save(outpath)


# ============ Excel(.xlsx) 输出：openpyxl 现绘，镜像官方版式 ============
_SONG = "宋体"
_THIN = Side(style="thin", color="000000")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_HEAD_FILL = PatternFill("solid", fgColor="D9E1F2")   # 表头浅蓝
_GROUP_FILL = PatternFill("solid", fgColor="F2F2F2")  # 分组小标题浅灰
# 表2「整改情况」/表3「处理情况」均为三档（客户口径 2026-07-16）——留给专家勾选
_T2_CHECK = "□符合　□基本符合　□不符合"
_T3_CHECK = "□符合　□基本符合　□不符合"


def _style(cell, *, bold=False, size=10, align="left", valign="center", wrap=True, fill=None):
    cell.font = Font(name=_SONG, size=size, bold=bold)
    cell.alignment = Alignment(horizontal=align, vertical=valign, wrap_text=wrap)
    cell.border = _BORDER
    if fill is not None:
        cell.fill = fill


def _xlsx_scaffold(ws, title, headers, widths):
    """写标题行(整行合并)+表头行，设列宽。返回下一可写行号(3)。"""
    ncol = len(headers)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncol)
    _style(ws.cell(row=1, column=1, value=title), bold=True, size=14, align="center")
    ws.row_dimensions[1].height = 28
    for j, h in enumerate(headers, start=1):
        _style(ws.cell(row=2, column=j, value=h), bold=True, align="center", fill=_HEAD_FILL)
    for j, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(j)].width = w
    return 3


def _footer_row(ws, row, ncol):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncol)
    _style(ws.cell(row=row, column=1, value="填表人：______    日期：____年__月__日    备注："), align="left")


def build_table2_xlsx(records, outpath):
    wb = Workbook()
    ws = wb.active
    ws.title = "表2"
    headers = ["序号", "整改要求", "整改主体", "评估内容", "评估方式", "佐证材料", "整改情况", "说明"]
    widths = [6, 34, 16, 26, 20, 30, 18, 14]
    r = _xlsx_scaffold(ws, "表2 事故整改和防范措施落实情况评估表", headers, widths)
    for rec in records:
        # 章节标题行((一)(二)(三))：主体列为"-"，整改情况也填"-"、不显示勾选项
        is_hdr = rec.get("subject", "").strip() == "-"
        vals = [rec.get("no", ""), rec.get("requirement", ""), rec.get("subject", ""),
                rec.get("content", ""), rec.get("method", ""), rec.get("materials", ""),
                ("-" if is_hdr else _T2_CHECK), rec.get("note", "")]
        for j, v in enumerate(vals, start=1):
            _style(ws.cell(row=r, column=j, value=v), bold=is_hdr,
                   align=("center" if j in (1, 7) else "left"), valign="top",
                   fill=(_GROUP_FILL if is_hdr else None))
        r += 1
    _footer_row(ws, r, len(headers))
    wb.save(outpath)


def build_table3_xlsx(groups, outpath):
    wb = Workbook()
    ws = wb.active
    ws.title = "表3"
    headers = ["序号", "事故责任人员/单位（括号内为隶属单位）", "处理建议", "提供材料单位", "佐证材料", "处理情况（是否与处理建议相符）", "说明"]
    widths = [6, 30, 30, 18, 30, 22, 14]
    ncol = len(headers)
    r = _xlsx_scaffold(ws, "表3 事故处理落实情况评估表", headers, widths)
    for g in groups:
        gt = g.get("group_title")
        if gt:                                    # 分组小标题：整行合并
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncol)
            _style(ws.cell(row=r, column=1, value=gt), bold=True, align="left", fill=_GROUP_FILL)
            r += 1
        for rec in g.get("rows", []):
            # 组标题行(如"行刑衔接情况")：其余列留空、不显示勾选项、加粗
            is_hdr = not (rec.get("suggestion") or rec.get("provider") or rec.get("materials"))
            vals = [rec.get("no", ""), rec.get("person", ""), rec.get("suggestion", ""),
                    rec.get("provider", ""), rec.get("materials", ""),
                    ("" if is_hdr else _T3_CHECK), rec.get("note", "")]
            for j, v in enumerate(vals, start=1):
                _style(ws.cell(row=r, column=j, value=v), bold=is_hdr,
                       align=("center" if j in (1, 6) else "left"), valign="top",
                       fill=(_GROUP_FILL if is_hdr else None))
            r += 1
    _footer_row(ws, r, ncol)
    wb.save(outpath)


def main():
    if len(sys.argv) < 3:
        print("用法: python generate_tables.py <input.json> <output_dir>")
        sys.exit(1)
    inp, outdir = sys.argv[1], sys.argv[2]
    os.makedirs(outdir, exist_ok=True)
    data = load_json_or_explain(inp)

    t2 = data.get("table2", [])
    t3 = data.get("table3_groups", [])
    out2 = os.path.join(outdir, "表2_事故整改和防范措施落实情况评估表.docx")
    out3 = os.path.join(outdir, "表3_事故处理落实情况评估表.docx")
    out2x = os.path.join(outdir, "表2_事故整改和防范措施落实情况评估表.xlsx")
    out3x = os.path.join(outdir, "表3_事故处理落实情况评估表.xlsx")
    build_table2(t2, out2)
    build_table3(t3, out3)
    build_table2_xlsx(t2, out2x)
    build_table3_xlsx(t3, out3x)
    print("OK: 已生成 4 个文件（表2/表3 各 docx + xlsx）")
    for p in (out2, out3, out2x, out3x):
        print(p)


if __name__ == "__main__":
    main()
